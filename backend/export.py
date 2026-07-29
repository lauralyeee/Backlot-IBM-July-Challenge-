"""
Feature 2 -- Assets to Producible Output.

Compiles a writer-chosen subset of confirmed canon assets into a document:
Character Bios, Location Breakdown, Beat Sheet, Pitch Packet, or Sample
Scene. Opposite direction to ingestion.py (assets -> document, not text ->
assets). Kept in its own module so main.py stays a single import + route
registration.

compile_document() produces Markdown (characters/locations/beats/pitch, a
narrow # / ## / "- " / "> " / **bold** dialect) or Fountain plain text
(script). render_download() then turns that into a PDF/DOCX/.fountain file
-- offline, deterministic, no LLM call, even on a re-download.

DOC_TYPES: str value = type_filter for db.list_assets(); None = all types.
"""

from __future__ import annotations

import io
import re
from datetime import date

# Maps the docType key (as sent by the client) to the asset `type` value it
# pulls from the database.  None means "pull all types" — same semantics as
# db.list_assets(type_filter=None).
DOC_TYPES: dict[str, str | None] = {
    "characters": "character",
    "locations": "location",
    "beats": None,
    "pitch": None,
    "script": None,
}

# Display labels chosen to match how these documents are actually named in
# the film/TV industry (cross-checked against real shooting scripts,
# schedules, and a character-bios one-sheet), not generic app-speak.
_DOC_LABELS: dict[str, str] = {
    "characters": "Character Bios",
    "locations": "Location Breakdown",
    "beats": "Beat Sheet",
    "pitch": "Pitch Packet",
    "script": "Sample Scene",
}

# Sampling caps for "pitch" — keeps the prompt tight and the output selective.
_PITCH_CAP: dict[str, int] = {
    "character": 8,
    "location": 6,
    "other": 999,     # factions/clans/etc. -- typically few, take all
    "event": 4,
}

# Sampling caps for "script" — a sample scene only needs a small cast and a
# single location, not the whole world; keeps the prompt tight and stops an
# large world from producing an unfocused, overcrowded scene.
_SCRIPT_CAP: dict[str, int] = {
    "character": 3,
    "location": 1,
    "event": 2,
    "other": 1,
}


def _today_str() -> str:
    """Portable 'Month D, YYYY' string (no %-d / %#d, which aren't
    cross-platform in strftime)."""
    d = date.today()
    return f"{d:%B} {d.day}, {d.year}"


# ── Era-order sort helper ─────────────────────────────────────────────────────

def _sort_by_era(assets: list[dict], world: dict) -> list[dict]:
    """Sort assets by their position in world["eras"] (the world's canonical
    chronological sequence), then preserve creation order within each era.
    Assets whose era isn't in world["eras"] sort last — they never crash."""
    eras = world.get("eras", [])
    return sorted(
        assets,
        key=lambda a: eras.index(a["era"]) if a.get("era") in eras else len(eras),
    )


# ── Sampling helpers ──────────────────────────────────────────────────────────

def _sample_by_cap(assets: list[dict], caps: dict[str, int]) -> list[dict]:
    """Shared by _pitch_sample and _script_sample: apply per-type sampling
    caps, in order, excluding any type not present in `caps` at all."""
    counts: dict[str, int] = {}
    out: list[dict] = []
    for a in assets:
        t = a.get("type", "")
        cap = caps.get(t)
        if cap is None:
            continue
        n = counts.get(t, 0)
        if n < cap:
            out.append(a)
            counts[t] = n + 1
    return out


def _pitch_sample(assets: list[dict]) -> list[dict]:
    """Apply per-type sampling caps for the Pitch Packet doc type.
    Any asset type not listed in _PITCH_CAP is excluded from the pitch entirely
    (e.g. raw lore dumps would bloat a pitch packet). Maintains order."""
    return _sample_by_cap(assets, _PITCH_CAP)


def _script_sample(assets: list[dict]) -> list[dict]:
    """Apply per-type sampling caps for the Sample Scene doc type — a small
    cast + one location, not the whole world. Maintains order."""
    return _sample_by_cap(assets, _SCRIPT_CAP)


# ── Prompt builders ──────────────────────────────────────────────────────────

def compile_system_prompt(doc_type: str, world: dict) -> str:
    label = _DOC_LABELS.get(doc_type, doc_type)
    if doc_type == "beats":
        return (
            f'You are a story editor and script supervisor for the world "{world["name"]}". '
            f"Your task is to compile a {label} in Markdown that reassembles the provided "
            "canon entries into narrative order, moving through the world's eras in sequence. "
            "Event entries are the narrative backbone — use them as beat headings. "
            "Character, location, and lore entries are supporting context — weave them in "
            "under the beats where relevant. "
            "Do NOT invent plot points, scenes, or story developments not evidenced by the "
            "given canon entries. Only sequence and connect what is already there. "
            "Use # for era headings, ## for individual beats/events, and bullet points for "
            "supporting context. Output Markdown only — no preamble, no closing remarks."
        )
    if doc_type == "pitch":
        return (
            f"You are a writer's room executive putting together a pitch packet for "
            f'"{world["name"]}" aimed at a producer or collaborator audience. '
            "Write in a punchy, evocative tone — hooks and atmosphere over exhaustive detail. "
            "Structure the document as: a one-paragraph logline/world overview, then brief "
            "## sections for Key Characters, Key Locations, and Story Hooks (drawn from event "
            "entries if present). "
            "Only use information present in the provided entries — do not invent details. "
            "Keep it tight: this is a room-leaving document, not a bible. Output Markdown only."
        )
    if doc_type == "script":
        return (
            f'You are a professional screenwriter drafting a sample scene set in the world "{world["name"]}". '
            "Write exactly ONE self-contained scene (roughly 1-3 pages), built only from the canon "
            "entries provided below — do not invent characters, locations, or plot points that "
            "aren't evidenced by those entries. This is a proof-of-concept scene meant to show what "
            "an actual script from this world could look like, not a finished draft.\n\n"
            "Output valid Fountain-syntax plain text (the industry-standard plain-text screenplay "
            "format — NOT Markdown, no #, no **, no bullet dashes). Follow this shape exactly:\n"
            "1. A title-page block at the very top, one 'Key: value' pair per line, then one blank "
            "   line before the scene itself:\n"
            "   Title: <a short scene/world title, in caps>\n"
            "   Credit: Written by\n"
            "   Author: World Copilot, drafted from canon\n"
            f"   Draft date: {_today_str()}\n"
            "2. A scene heading on its own line, in CAPS, starting with INT. or EXT., e.g. "
            "   'INT. THE SUNKEN LIBRARY - NIGHT'.\n"
            "3. Action lines directly under it: plain prose, present tense, no dialogue.\n"
            "4. A character cue: that character's name ALONE on its own line, in CAPS, with a "
            "   blank line before and after it.\n"
            "5. Their dialogue: plain text on the line(s) immediately following the cue.\n"
            "6. A parenthetical (only if it earns its place): a short direction in parentheses, "
            "   on its own line, between a character cue and their dialogue.\n"
            "7. A transition (optional, sparingly): CAPS, ending in 'TO:', e.g. 'CUT TO:', on its "
            "   own line.\n"
            "Only use character names and locations that appear in the canon entries below. "
            "Output the Fountain text only — no preamble, no explanation, no commentary before or "
            "after it."
        )
    # characters — prose bios (matches how a real production one-sheet lists
    # a cast: a bolded name lead-in, then flowing sentences, not bullets)
    if doc_type == "characters":
        return (
            f'You are a script supervisor and production coordinator for the world "{world["name"]}". '
            "Compile a Character Bios document from the canon entries provided, in the voice of a "
            "real production one-sheet: for each character, start with their name in bold as a "
            "lead-in (e.g. **NAME** is ...), then 2-4 flowing sentences of prose covering age (if "
            "known), background, personality, and voice/dialect notes where present. One paragraph "
            "per character, separated by a blank line — no bullet points, no sub-headers per entry. "
            "Only use information present in the provided entries — do not invent details not given. "
            "Do not add commentary, preamble, or closing remarks. Output Markdown only (plain "
            "paragraphs with a **bold** name lead-in), no # or ## headers."
        )
    # locations — original behaviour unchanged (era-grouped headers + bullets
    # reads like a real AD location breakdown, which is naturally list-like)
    return (
        f'You are a script supervisor and production coordinator for the world "{world["name"]}". '
        f"Your task is to compile a {label} in Markdown from the canon entries provided. "
        "Format the document with a clear Markdown header (##) for each entry. "
        "Under each header include the relevant production details as a short bulleted list. "
        "Only use information present in the provided entries — do not invent details not given. "
        "Do not add commentary, preamble, or closing remarks. Output Markdown only."
    )


def compile_user_prompt(doc_type: str, assets: list[dict], world: dict) -> str:
    if doc_type == "beats":
        sorted_assets = _sort_by_era(assets, world)
        eras = world.get("eras", [])
        has_events = any(a.get("type") == "event" for a in sorted_assets)
        title_line = _title_line(doc_type, world, len(sorted_assets))
        no_events_note = (
            "" if has_events else
            "\n\nNote: no dedicated 'event' assets exist yet in this world — "
            "produce an era-by-era outline from the available characters, locations, "
            "and lore entries instead.\n"
        )
        era_order_note = (
            f"Era order for this world (chronological): {' → '.join(eras)}.\n"
            if eras else ""
        )
        instruction = (
            f"Produce a Beat Sheet. {era_order_note}"
            "Organise the entries below into narrative order by era, using event entries "
            "as beat headings and weaving character/location/lore entries in as context. "
            "Do not invent plot points not evidenced by the canon entries."
            f"{no_events_note}"
        )
        entries = "\n\n".join(_format_asset(a) for a in sorted_assets)
        return f"{instruction}\n\nDocument title: {title_line}\n\nCANON ENTRIES:\n\n{entries}"

    if doc_type == "pitch":
        sampled = _pitch_sample(assets)
        title_line = _title_line(doc_type, world, len(sampled))
        instruction = (
            "Produce a Pitch Packet. Write a tight, evocative overview for a producer audience: "
            "one-paragraph logline/world overview, then ## sections for Key Characters, "
            "Key Locations, and Story Hooks. Prioritise atmosphere and hooks over exhaustive "
            "detail. Only use information present in the entries below."
        )
        entries = "\n\n".join(_format_asset(a) for a in sampled)
        return f"{instruction}\n\nDocument title: {title_line}\n\nCANON ENTRIES (selected):\n\n{entries}"

    if doc_type == "script":
        sampled = _script_sample(assets)
        instruction = (
            "Produce a Sample Scene in Fountain syntax, exactly as specified in your instructions "
            "above, using only the characters/locations/events below. Pick whichever subset best "
            "supports one coherent scene — you do not need to use every entry."
        )
        entries = "\n\n".join(_format_asset(a) for a in sampled)
        return f"{instruction}\n\nCANON ENTRIES (selected):\n\n{entries}"

    # characters / locations — original behaviour unchanged
    title_line = _title_line(doc_type, world, len(assets))
    if doc_type == "characters":
        instruction = (
            "Produce a Character Bios document. For each character entry below, write a bio "
            "paragraph in the style described above (bold name lead-in, then flowing prose) "
            "covering their description/backstory, voice or dialect notes if present, era, and "
            "faction where relevant. Keep each bio concise and production-ready."
        )
    else:  # locations
        instruction = (
            "Produce a Location Breakdown grouped by era. For each era, write a # heading with the era name, "
            "then a ## heading for each location in that era, with bullet points covering: "
            "description, and faction/controlling faction if present. "
            "Keep each entry concise and production-ready."
        )
    entries = "\n\n".join(_format_asset(a) for a in assets)
    return f"{instruction}\n\nDocument title: {title_line}\n\nCANON ENTRIES:\n\n{entries}"


def _format_asset(asset: dict) -> str:
    """Serialize one asset for the prompt — same fields as retrieval.canon_block
    but with a 400-char content cap (larger than canon_block's 220 chars since
    this output is meant to be read in full, not just used as grounding context)."""
    content = (asset.get("content") or "").strip()
    if len(content) > 400:
        content = content[:400].rsplit(" ", 1)[0] + "…"
    faction = asset.get("faction") or "—"
    return (
        f"Title: {asset.get('title', 'Untitled')}\n"
        f"Type: {asset.get('type', '')}\n"
        f"Era: {asset.get('era', '')}\n"
        f"Faction: {faction}\n"
        f"Content: {content}"
    )


# ── Title line helper ────────────────────────────────────────────────────────

def _title_line(doc_type: str, world: dict, count: int) -> str:
    label = _DOC_LABELS.get(doc_type, doc_type)
    return f"# {world['name']}: {label} ({count} {'entry' if count == 1 else 'entries'})"


# ── Offline fallback ─────────────────────────────────────────────────────────

def _offline_era_group(assets: list[dict], world: dict, include_type: bool = False) -> list[str]:
    """Shared helper: group assets by era in world order, return Markdown lines.
    Used by both the 'locations' and 'beats' offline branches."""
    eras = world.get("eras", [])
    sorted_assets = _sort_by_era(assets, world)
    # Build ordered era list preserving world era order, then appending any
    # unknown eras that appear in the data.
    seen_eras: list[str] = []
    for a in sorted_assets:
        era_key = (a.get("era") or "Unspecified").strip()
        if era_key not in seen_eras:
            seen_eras.append(era_key)

    by_era: dict[str, list[dict]] = {}
    for a in sorted_assets:
        era_key = (a.get("era") or "Unspecified").strip()
        by_era.setdefault(era_key, []).append(a)

    lines: list[str] = []
    for era_key in seen_eras:
        lines.append(f"# {era_key}")
        lines.append("")
        for a in by_era[era_key]:
            lines.append(f"## {a.get('title', 'Untitled')}")
            if include_type:
                lines.append(f"- **Type:** {a.get('type', '—')}")
            content = (a.get("content") or "").strip()
            if content:
                lines.append(f"- {content[:400]}")
            faction = (a.get("faction") or "—").strip()
            if faction and faction != "—":
                lines.append(f"- **Faction:** {faction}")
            lines.append("")
    return lines


def _offline_script(assets: list[dict], world: dict) -> str:
    """Deterministic, no-LLM fallback for 'script': a minimal but real
    Fountain-syntax scene template built directly from asset fields — a
    location's description as an establishing action line, then a short
    two-line exchange templated from up to two characters' own content."""
    sampled = _script_sample(assets)
    chars = [a for a in sampled if a.get("type") == "character"][:2]
    locs = [a for a in sampled if a.get("type") == "location"][:1]
    loc_title = (locs[0].get("title") if locs else "AN UNSPECIFIED LOCATION") or "AN UNSPECIFIED LOCATION"

    lines = [
        f"Title: {world['name'].upper()}",
        "Credit: Written by",
        "Author: World Copilot (offline draft, watsonx was unreachable)",
        f"Draft date: {_today_str()}",
        "",
        f"INT. {loc_title.upper()} - DAY",
        "",
    ]
    if locs and (locs[0].get("content") or "").strip():
        content = locs[0]["content"].strip()
        lines.append(content[:300] + ("…" if len(content) > 300 else ""))
        lines.append("")
    else:
        lines.append("The space sits mid-scene, waiting for the story to catch up to it.")
        lines.append("")

    for c in chars:
        lines.append((c.get("title") or "CHARACTER").upper())
        lines.append("")
        content = (c.get("content") or "").strip()
        snippet = content[:140] + "…" if len(content) > 140 else content
        lines.append(snippet or "(A line reflecting this character's voice goes here.)")
        lines.append("")

    if not chars:
        lines.append("Note: no confirmed character entries were available for this offline draft.")
        lines.append("")

    return "\n".join(lines)


def offline_compile(doc_type: str, assets: list[dict], world: dict) -> str:
    """Deterministic, no-LLM fallback: builds the same document directly from
    asset fields using plain Markdown headers and bullets (or, for 'script',
    plain Fountain syntax). Returns a str — the caller can drop it straight
    into the response without further processing."""
    if doc_type == "script":
        # Fountain, not Markdown — no title-line/warning-blockquote wrapper,
        # since a real screenplay's title page IS the document header.
        return _offline_script(assets, world)

    lines: list[str] = [_title_line(doc_type, world, len(assets))]
    lines.append("")
    lines.append("> ⚠ Generated offline, watsonx was unreachable. Structure built directly from canon fields.")
    lines.append("")

    if doc_type == "locations":
        lines.extend(_offline_era_group(assets, world, include_type=False))

    elif doc_type == "beats":
        # Mixed-type view — include type label so the writer can tell entries apart
        lines.extend(_offline_era_group(assets, world, include_type=True))

    elif doc_type == "pitch":
        # Apply the same sampling cap as the online path, then group by type
        sampled = _pitch_sample(assets)
        # Group by type for readability
        by_type: dict[str, list[dict]] = {}
        for a in sampled:
            by_type.setdefault(a.get("type", "other"), []).append(a)
        type_order = ["character", "location", "event", "other"]
        for t in type_order:
            if t not in by_type:
                continue
            lines.append(f"# {t.capitalize()}s")
            lines.append("")
            for a in by_type[t]:
                lines.append(f"## {a.get('title', 'Untitled')}")
                content = (a.get("content") or "").strip()
                if content:
                    lines.append(f"- {content[:400]}")
                era = (a.get("era") or "").strip()
                if era:
                    lines.append(f"- **Era:** {era}")
                faction = (a.get("faction") or "—").strip()
                if faction and faction != "—":
                    lines.append(f"- **Faction:** {faction}")
                lines.append("")

    else:
        # characters — one prose bio per entry (bold name lead-in, matching
        # the online path's real-one-sheet style instead of a bullet list)
        for a in assets:
            title = a.get("title", "Untitled")
            content = (a.get("content") or "").strip()
            era = (a.get("era") or "").strip()
            faction = (a.get("faction") or "").strip()
            bits = [content] if content else []
            if era:
                bits.append(f"Era: {era}.")
            if faction and faction != "—":
                bits.append(f"Affiliated with {faction}.")
            body = " ".join(b.rstrip(".") + "." for b in bits if b) or "No further details recorded yet."
            lines.append(f"**{title}** {body}")
            lines.append("")

    return "\n".join(lines)


# ── Format conversion (Markdown/Fountain text → PDF / DOCX / Fountain file) ──
# Offline and deterministic -- no LLM call, no DB access. Converts whatever
# text the compile step already produced into a downloadable file. Two
# narrow parsers do the work: _md_blocks() only understands the Markdown
# subset this module emits, _parse_fountain() only the Fountain subset the
# "script" prompt asks for. Neither is a general-purpose parser.

_TITLE_KEYS = {"title", "credit", "author", "authors", "draft date", "source", "contact"}

# Matches 1-6 leading '#' characters. Granite is only asked for "#"/"##" but
# sometimes emits a deeper "### " heading anyway; without this it fell through
# to the plain-paragraph branch and a literal "###" leaked into exported
# files. Any heading depth now renders as a real heading, capped at h3.
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def _md_blocks(markdown_text: str) -> list[dict]:
    """Turn the Markdown this module emits into a small block list:
    {"type": "h1"|"h2"|"h3"|"li"|"quote"|"p", "text": ...}."""
    blocks: list[dict] = []
    para_buf: list[str] = []

    def flush_para():
        if para_buf:
            blocks.append({"type": "p", "text": " ".join(para_buf).strip()})
            para_buf.clear()

    for raw in (markdown_text or "").split("\n"):
        stripped = raw.strip()
        heading = _HEADING_RE.match(stripped) if stripped else None
        if not stripped:
            flush_para()
        elif heading:
            flush_para()
            depth = len(heading.group(1))
            kind = "h1" if depth == 1 else "h2" if depth == 2 else "h3"
            blocks.append({"type": kind, "text": heading.group(2).strip()})
        elif stripped.startswith("> "):
            flush_para()
            blocks.append({"type": "quote", "text": stripped[2:].strip()})
        elif stripped.startswith("- "):
            flush_para()
            blocks.append({"type": "li", "text": stripped[2:].strip()})
        else:
            para_buf.append(stripped)
    flush_para()
    return blocks


def _inline_to_reportlab(text: str) -> str:
    """**bold** -> <b>bold</b>; escapes stray angle brackets/amps first so
    reportlab's mini-HTML Paragraph markup doesn't choke on canon text that
    happens to contain a literal '<' or '&'."""
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)


def _parse_fountain(text: str) -> tuple[dict, list[dict]]:
    """Small Fountain-syntax parser covering exactly the subset this module
    asks Granite (or the offline fallback) to produce: an optional title-page
    key:value block, then scene headings, action, character cues + optional
    parenthetical + dialogue, and transitions."""
    lines = [l.rstrip() for l in (text or "").split("\n")]
    i = 0
    n = len(lines)
    title_page: dict[str, str] = {}

    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1
            break
        m = re.match(r"^([A-Za-z ]+):\s*(.*)$", line)
        if m and m.group(1).strip().lower() in _TITLE_KEYS:
            title_page[m.group(1).strip().lower()] = m.group(2).strip()
            i += 1
        else:
            break

    def is_cue_line(s: str) -> bool:
        return bool(s) and s == s.upper() and any(c.isalpha() for c in s) and len(s) < 40

    blocks: list[dict] = []
    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        upper = line.upper()

        if re.match(r"^(INT|EXT|EST|I/E)[./\s]", upper) or upper.startswith("INT/EXT"):
            blocks.append({"type": "heading", "text": upper})
            i += 1
        elif is_cue_line(line) and upper.rstrip(":").endswith("TO"):
            blocks.append({"type": "transition", "text": upper})
            i += 1
        elif is_cue_line(line) and i + 1 < n and lines[i + 1].strip() != "":
            name = upper.rstrip(":")
            i += 1
            paren = ""
            if i < n and lines[i].strip().startswith("(") and lines[i].strip().endswith(")"):
                paren = lines[i].strip()
                i += 1
            dialogue_lines: list[str] = []
            while i < n and lines[i].strip() and not is_cue_line(lines[i].strip()):
                dialogue_lines.append(lines[i].strip())
                i += 1
            blocks.append({
                "type": "dialogue",
                "character": name,
                "parenthetical": paren,
                "text": " ".join(dialogue_lines).strip(),
            })
        else:
            action_lines = [line]
            i += 1
            while i < n and lines[i].strip() and not is_cue_line(lines[i].strip()):
                action_lines.append(lines[i].strip())
                i += 1
            blocks.append({"type": "action", "text": " ".join(action_lines).strip()})

    return title_page, blocks


def markdown_to_pdf(markdown_text: str, *, doc_title: str, doc_subtitle: str = "") -> bytes:
    """Renders a Markdown-compiled document (characters/locations/beats/pitch)
    as a clean, modern PDF: a title block, running footer with page numbers,
    and styled headings/bullets/paragraphs — built with reportlab's Platypus
    flowables so pagination/wrapping is handled for us."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, ListFlowable, ListItem, HRFlowable

    buf = io.BytesIO()

    styles = {
        "title": ParagraphStyle("title", fontName="Helvetica-Bold", fontSize=22, leading=26, spaceAfter=4, textColor=colors.HexColor("#141414")),
        "subtitle": ParagraphStyle("subtitle", fontName="Helvetica", fontSize=10.5, leading=14, textColor=colors.HexColor("#6B6B6A"), spaceAfter=18),
        "h1": ParagraphStyle("h1", fontName="Helvetica-Bold", fontSize=15, leading=19, spaceBefore=18, spaceAfter=8, textColor=colors.HexColor("#141414")),
        "h2": ParagraphStyle("h2", fontName="Helvetica-Bold", fontSize=12.5, leading=16, spaceBefore=14, spaceAfter=6, textColor=colors.HexColor("#141414")),
        "h3": ParagraphStyle("h3", fontName="Helvetica-Bold", fontSize=11, leading=14.5, spaceBefore=11, spaceAfter=5, textColor=colors.HexColor("#141414")),
        "body": ParagraphStyle("body", fontName="Times-Roman", fontSize=10.5, leading=15.5, spaceAfter=9, textColor=colors.HexColor("#1C1C1B")),
        "li": ParagraphStyle("li", fontName="Times-Roman", fontSize=10.5, leading=15, spaceAfter=3, leftIndent=4, textColor=colors.HexColor("#1C1C1B")),
        "quote": ParagraphStyle("quote", fontName="Helvetica-Oblique", fontSize=9.5, leading=13, textColor=colors.HexColor("#8A4A42"), spaceAfter=10),
    }

    def _footer(c, doc):
        c.saveState()
        c.setFont("Helvetica", 8)
        c.setFillColor(colors.HexColor("#9A9A99"))
        c.drawString(0.85 * inch, 0.55 * inch, doc_title)
        c.drawRightString(LETTER[0] - 0.85 * inch, 0.55 * inch, f"Page {doc.page}")
        c.restoreState()

    story = [Paragraph(_inline_to_reportlab(doc_title), styles["title"])]
    if doc_subtitle:
        story.append(Paragraph(_inline_to_reportlab(doc_subtitle), styles["subtitle"]))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#D6D5D4"), spaceAfter=14))

    li_buf: list[str] = []

    def flush_li():
        nonlocal li_buf
        if li_buf:
            story.append(ListFlowable(
                [ListItem(Paragraph(_inline_to_reportlab(t), styles["li"]), bulletColor=colors.HexColor("#6B6B6A")) for t in li_buf],
                bulletType="bullet", leftIndent=16, bulletFontSize=7.5,
            ))
            li_buf = []

    for b in _md_blocks(markdown_text):
        if b["type"] == "li":
            li_buf.append(b["text"])
            continue
        flush_li()
        if b["type"] == "h1":
            story.append(Paragraph(_inline_to_reportlab(b["text"]), styles["h1"]))
        elif b["type"] == "h2":
            story.append(Paragraph(_inline_to_reportlab(b["text"]), styles["h2"]))
        elif b["type"] == "h3":
            story.append(Paragraph(_inline_to_reportlab(b["text"]), styles["h3"]))
        elif b["type"] == "quote":
            story.append(Paragraph(_inline_to_reportlab(b["text"]), styles["quote"]))
        elif b["type"] == "p":
            story.append(Paragraph(_inline_to_reportlab(b["text"]), styles["body"]))
    flush_li()

    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.9 * inch, rightMargin=0.9 * inch, topMargin=0.9 * inch, bottomMargin=0.9 * inch,
        title=doc_title,
    )
    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


def markdown_to_docx(markdown_text: str, *, doc_title: str, doc_subtitle: str = "") -> bytes:
    """Renders a Markdown-compiled document as an editable Word file — the
    format most producers/executives actually mark up and send back."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor

    doc = Document()
    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(0.9)
        section.top_margin = section.bottom_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.5)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(doc_title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = "Calibri"

    if doc_subtitle:
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(doc_subtitle)
        sub_run.font.size = Pt(10)
        sub_run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6A)

    doc.add_paragraph()

    def add_runs(paragraph, text: str):
        for part in re.split(r"(\*\*.+?\*\*)", text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = paragraph.add_run(part[2:-2])
                r.bold = True
            else:
                paragraph.add_run(part)

    for b in _md_blocks(markdown_text):
        if b["type"] == "h1":
            add_runs(doc.add_heading(level=1), b["text"])
        elif b["type"] == "h2":
            add_runs(doc.add_heading(level=2), b["text"])
        elif b["type"] == "h3":
            add_runs(doc.add_heading(level=3), b["text"])
        elif b["type"] == "li":
            add_runs(doc.add_paragraph(style="List Bullet"), b["text"])
        elif b["type"] == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r = p.add_run(b["text"])
            r.italic = True
            r.font.color.rgb = RGBColor(0x8A, 0x4A, 0x42)
        elif b["type"] == "p":
            add_runs(doc.add_paragraph(), b["text"])

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def fountain_to_pdf(fountain_text: str) -> bytes:
    """Renders Fountain-syntax scene text as a real, industry-formatted
    screenplay PDF: Courier 12, a title page, 1.5in left / 1in right margin,
    centered character cues/dialogue column, right-set transitions."""
    import textwrap
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas as _canvas

    title_page, blocks = _parse_fountain(fountain_text)

    buf = io.BytesIO()
    c = _canvas.Canvas(buf, pagesize=LETTER)
    width, height = LETTER

    LEFT = 1.5 * inch
    RIGHT = width - 1.0 * inch
    TOP = height - 1.0 * inch
    BOTTOM = 1.0 * inch
    FONT = "Courier"
    SIZE = 12
    LEAD = 14.4

    page_num = [1]
    y = [TOP]

    def wrap(text: str, max_chars: int) -> list[str]:
        return textwrap.wrap(text, max_chars) or [""]

    def new_page():
        c.showPage()
        page_num[0] += 1
        y[0] = TOP
        c.setFont(FONT, SIZE)
        c.setFillColor(colors.black)

    def ensure(space: float = LEAD):
        if y[0] - space < BOTTOM:
            new_page()

    # ---- Title page ----
    title = title_page.get("title", "UNTITLED")
    c.setFont("Courier-Bold", 15)
    c.setFillColor(colors.black)
    c.drawCentredString(width / 2, height * 0.44, title)
    c.setFont(FONT, 12)
    y_tp = height * 0.44 - 30
    for key in ("credit", "author", "authors"):
        if title_page.get(key):
            c.drawCentredString(width / 2, y_tp, title_page[key])
            y_tp -= 18
    if title_page.get("draft date"):
        c.drawString(LEFT, BOTTOM + 4, title_page["draft date"])
    c.setFont("Courier-Oblique", 9)
    c.setFillColor(colors.HexColor("#8A8A89"))
    c.drawCentredString(width / 2, BOTTOM + 24,
                         "Sample scene, generated from world canon (proof of concept, not a shooting draft)")
    c.showPage()

    c.setFont(FONT, SIZE)
    c.setFillColor(colors.black)
    y[0] = TOP

    for b in blocks:
        if b["type"] == "heading":
            ensure(LEAD * 2)
            c.setFont("Courier-Bold", SIZE)
            c.drawString(LEFT, y[0], b["text"])
            c.setFont(FONT, SIZE)
            y[0] -= LEAD * 1.7
        elif b["type"] == "action":
            for line in wrap(b["text"], 60):
                ensure()
                c.drawString(LEFT, y[0], line)
                y[0] -= LEAD
            y[0] -= LEAD * 0.5
        elif b["type"] == "transition":
            ensure(LEAD * 1.6)
            c.drawRightString(RIGHT, y[0], b["text"])
            y[0] -= LEAD * 1.6
        elif b["type"] == "dialogue":
            ensure(LEAD * 1.6)
            cue_x = width / 2 - 0.9 * inch
            c.drawString(cue_x, y[0], b["character"])
            y[0] -= LEAD
            if b.get("parenthetical"):
                c.setFont("Courier-Oblique", SIZE - 1)
                c.drawString(cue_x + 0.25 * inch, y[0], b["parenthetical"])
                c.setFont(FONT, SIZE)
                y[0] -= LEAD
            for line in wrap(b["text"], 33):
                ensure()
                c.drawString(cue_x - 0.35 * inch, y[0], line)
                y[0] -= LEAD
            y[0] -= LEAD * 0.5

    c.setFont(FONT, 9)
    c.setFillColor(colors.HexColor("#9A9A99"))
    c.drawRightString(RIGHT, height - 0.6 * inch, f"{page_num[0]}.")
    c.save()
    return buf.getvalue()


def fountain_to_docx(fountain_text: str) -> bytes:
    """Renders Fountain-syntax scene text as a Word doc, screenplay-styled
    (Courier New, centered/indented dialogue column) rather than a literal
    text dump — importable/editable, and print-recognizable as a script."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title_page, blocks = _parse_fountain(fountain_text)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Courier New"
    normal.font.size = Pt(12)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(title_page.get("title", "UNTITLED"))
    r.bold = True
    for key in ("credit", "author", "authors"):
        if title_page.get(key):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(title_page[key])
    if title_page.get("draft date"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(title_page["draft date"])
    doc.add_page_break()

    for b in blocks:
        if b["type"] == "heading":
            p = doc.add_paragraph()
            r = p.add_run(b["text"])
            r.bold = True
        elif b["type"] == "action":
            doc.add_paragraph(b["text"])
        elif b["type"] == "transition":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(b["text"])
        elif b["type"] == "dialogue":
            cue_p = doc.add_paragraph()
            cue_p.paragraph_format.left_indent = Inches(1.6)
            cue_p.add_run(b["character"])
            if b.get("parenthetical"):
                par_p = doc.add_paragraph()
                par_p.paragraph_format.left_indent = Inches(1.4)
                par_run = par_p.add_run(b["parenthetical"])
                par_run.italic = True
            d_p = doc.add_paragraph(b["text"])
            d_p.paragraph_format.left_indent = Inches(1.0)
            d_p.paragraph_format.right_indent = Inches(1.5)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def render_download(doc_type: str, content: str, fmt: str, *, world_name: str, asset_count: int = 0) -> tuple[bytes, str, str]:
    """Converts already-compiled text (Markdown for characters/locations/
    beats/pitch, Fountain for script) into the requested downloadable file.
    Never calls the LLM or the DB — pure, deterministic, offline conversion
    of text the caller already has. Returns (bytes, media_type, filename)."""
    label = _DOC_LABELS.get(doc_type, doc_type)
    safe_stub = re.sub(r"[^A-Za-z0-9]+", "-", f"{world_name}-{label}").strip("-").lower() or "export"

    if fmt == "markdown":
        return content.encode("utf-8"), "text/markdown; charset=utf-8", f"{safe_stub}.md"

    if doc_type == "script":
        if fmt == "fountain":
            return content.encode("utf-8"), "text/plain; charset=utf-8", f"{safe_stub}.fountain"
        if fmt == "pdf":
            return fountain_to_pdf(content), "application/pdf", f"{safe_stub}.pdf"
        if fmt == "docx":
            return (
                fountain_to_docx(content),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{safe_stub}.docx",
            )
    else:
        subtitle = f"{world_name} — {asset_count} {'entry' if asset_count == 1 else 'entries'} compiled"
        if fmt == "pdf":
            return markdown_to_pdf(content, doc_title=label, doc_subtitle=subtitle), "application/pdf", f"{safe_stub}.pdf"
        if fmt == "docx":
            return (
                markdown_to_docx(content, doc_title=label, doc_subtitle=subtitle),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{safe_stub}.docx",
            )

    raise ValueError(f"Unsupported format '{fmt}' for docType '{doc_type}'")
